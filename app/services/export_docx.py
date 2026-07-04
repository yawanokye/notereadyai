import re
import uuid
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it.token import Token

from app.services.markdown_render import markdown_parser


_BRAND = "3346A8"
_DARK = "172033"
_LIGHT = "EEF1FF"
_GREY = "667085"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(_GREY)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _configure_document(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(_DARK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.15

    heading_sizes = {1: 18, 2: 15, 3: 12.5, 4: 11.5}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(_BRAND if level <= 2 else _DARK)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
        style.paragraph_format.space_after = Pt(6)

    document.core_properties.title = title
    document.core_properties.author = "NoteReady AI"
    document.core_properties.subject = "Lecture notes and study material"

    header = section.header.paragraphs[0]
    header.text = "NoteReady AI"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.name = "Aptos"
        header.runs[0].font.size = Pt(8.5)
        header.runs[0].font.color.rgb = RGBColor.from_string(_GREY)

    footer = section.footer.paragraphs[0]
    _add_page_number(footer)


def _add_cover(document: Document, title: str, markdown_text: str) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(90)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor.from_string(_BRAND)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Web-ready lecture and study material")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(_GREY)

    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(28)
    brand_run = brand.add_run("Prepared with NoteReady AI")
    brand_run.bold = True
    brand_run.font.size = Pt(10)
    brand_run.font.color.rgb = RGBColor.from_string(_DARK)

    headings = re.findall(r"^#\s+(.+)$", markdown_text, flags=re.MULTILINE)
    if len(headings) > 1:
        document.add_page_break()
        document.add_heading("Contents", level=1)
        for heading in headings:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.12)
            run = paragraph.add_run(heading.strip())
            run.font.size = Pt(10.5)
        document.add_page_break()
    else:
        document.add_page_break()


def _append_inline(paragraph, inline: Token | None) -> None:
    if inline is None:
        return
    bold = False
    italic = False
    code = False
    link_depth = 0
    for child in inline.children or []:
        if child.type == "strong_open":
            bold = True
        elif child.type == "strong_close":
            bold = False
        elif child.type == "em_open":
            italic = True
        elif child.type == "em_close":
            italic = False
        elif child.type == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(_DARK)
        elif child.type == "link_open":
            link_depth += 1
        elif child.type == "link_close":
            link_depth = max(0, link_depth - 1)
        elif child.type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break()
        elif child.type == "image":
            run = paragraph.add_run(f"[Figure: {child.content or 'image'}]")
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(_GREY)
        elif child.type == "text":
            run = paragraph.add_run(child.content)
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = "Aptos Mono"
            if link_depth:
                run.underline = True
                run.font.color.rgb = RGBColor.from_string(_BRAND)


def _plain_inline(inline: Token | None) -> str:
    if inline is None:
        return ""
    parts: list[str] = []
    for child in inline.children or []:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append(" ")
        elif child.type == "image":
            parts.append(child.content or "image")
    return "".join(parts).strip()


def _add_body_paragraph(
    document: Document,
    inline: Token | None,
    *,
    prefix: str = "",
    indent_level: int = 0,
    quote: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    if prefix:
        prefix_run = paragraph.add_run(prefix)
        prefix_run.bold = True
    _append_inline(paragraph, inline)
    if indent_level:
        paragraph.paragraph_format.left_indent = Inches(0.24 * indent_level)
        paragraph.paragraph_format.first_line_indent = Inches(-0.14)
    if quote:
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(8)
        for run in paragraph.runs:
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(_GREY)


def _parse_table(tokens: list[Token], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    current_row: list[str] | None = None
    current_cell = ""
    index = start_index + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == "table_close":
            return rows, index
        if token.type == "tr_open":
            current_row = []
        elif token.type in {"th_open", "td_open"}:
            current_cell = ""
        elif token.type == "inline":
            current_cell = _plain_inline(token)
        elif token.type in {"th_close", "td_close"} and current_row is not None:
            current_row.append(current_cell)
        elif token.type == "tr_close" and current_row is not None:
            rows.append(current_row)
            current_row = None
        index += 1
    return rows, index


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        _set_row_cant_split(row)
        if row_index == 0:
            _set_repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value)
            run.font.name = "Aptos"
            run.font.size = Pt(9)
            if row_index == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(_DARK)
                _set_cell_shading(cell, _LIGHT)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _render_markdown(document: Document, markdown_text: str) -> None:
    tokens = markdown_parser().parse(markdown_text)
    list_stack: list[dict[str, int | str | bool]] = []
    quote_depth = 0
    h1_seen = False
    index = 0

    while index < len(tokens):
        token = tokens[index]

        if token.type == "table_open":
            rows, index = _parse_table(tokens, index)
            _add_table(document, rows)

        elif token.type == "heading_open":
            level = int(token.tag[1]) if token.tag.startswith("h") else 2
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            text = _plain_inline(inline)
            if level == 1:
                if h1_seen:
                    document.add_page_break()
                h1_seen = True
            heading = document.add_heading(level=min(level, 4))
            heading.clear()
            _append_inline(heading, inline)
            index += 2  # Skip inline and heading_close.

        elif token.type == "bullet_list_open":
            list_stack.append({"type": "bullet", "counter": 0, "first": True})

        elif token.type == "ordered_list_open":
            start = 1
            if token.attrs:
                try:
                    start = int(dict(token.attrs).get("start", 1))
                except (TypeError, ValueError):
                    start = 1
            list_stack.append({"type": "ordered", "counter": start - 1, "first": True})

        elif token.type in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()

        elif token.type == "list_item_open" and list_stack:
            list_stack[-1]["counter"] = int(list_stack[-1]["counter"]) + 1
            list_stack[-1]["first"] = True

        elif token.type == "blockquote_open":
            quote_depth += 1

        elif token.type == "blockquote_close":
            quote_depth = max(0, quote_depth - 1)

        elif token.type == "paragraph_open":
            inline = tokens[index + 1] if index + 1 < len(tokens) and tokens[index + 1].type == "inline" else None
            prefix = ""
            indent_level = 0
            if list_stack:
                current = list_stack[-1]
                if bool(current.get("first", True)):
                    if current["type"] == "ordered":
                        prefix = f"{current['counter']}. "
                    else:
                        prefix = "• "
                    current["first"] = False
                indent_level = len(list_stack)
            _add_body_paragraph(
                document,
                inline,
                prefix=prefix,
                indent_level=indent_level,
                quote=quote_depth > 0,
            )
            index += 2  # Skip inline and paragraph_close.

        elif token.type == "fence":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(token.content.rstrip())
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9)

        elif token.type == "code_block":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(token.content.rstrip())
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9)

        elif token.type == "hr":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "D7DCE5")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)

        index += 1


def build_docx(title: str, markdown_text: str, output_dir: Path) -> Path:
    document = Document()
    _configure_document(document, title)
    _add_cover(document, title, markdown_text)
    _render_markdown(document, markdown_text)

    filename = f"{uuid.uuid4().hex}_{_safe_filename(title)}.docx"
    path = output_dir / filename
    document.save(path)
    return path


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")
    return cleaned[:80] or "noteready_output"
