from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


def test_lecture_notes_batched_preview() -> None:
    client = TestClient(app)
    response = client.post(
        "/api/lecture-notes/jobs",
        data={
            "academic_level": "Undergraduate Level 100",
            "course_title": "Principles of Procurement",
            "course_code": "PCM 102",
            "credit_hours": "3",
            "teaching_weeks": "2",
            "contact_hours_per_week": "3",
            "citation_style": "APA 7th edition",
            "context_preference": "Use Ghanaian examples where relevant.",
        },
        files={
            "file": (
                "outline.txt",
                b"Week 1: Introduction to procurement\nMeaning and scope\n\nWeek 2: Supplier selection\nEvaluation criteria",
                "text/plain",
            )
        },
    )
    assert response.status_code == 200
    job = response.json()
    assert job["title"] == "Principles of Procurement Lecture Notes"
    assert job["ai_enabled"] is False
    assert len(job["modules"]) == 2

    module_id = job["modules"][0]["id"]
    batch_response = client.post(
        f"/api/lecture-notes/jobs/{job['job_id']}/batches/{module_id}"
    )
    assert batch_response.status_code == 200
    batch = batch_response.json()
    assert "Development Preview" in batch["content_markdown"]
    assert "<h1>Development Preview</h1>" in batch["content_html"]


def test_docx_export_renders_markdown_table_and_inline_formatting() -> None:
    client = TestClient(app)
    markdown = (
        "# Topic\n\nThis is **important**.\n\n"
        "| Item | Meaning |\n|---|---|\n| A | First |\n| B | Second |\n\n"
        "1. Question one\n2. Question two"
    )
    response = client.post(
        "/api/exports/docx",
        json={"title": "Sample Notes", "content_markdown": markdown},
    )
    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Item"
    assert any(run.bold and "important" in run.text for p in document.paragraphs for run in p.runs)
